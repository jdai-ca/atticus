"""
Convert Jupyter Notebooks (.ipynb) to formatted, academic-grade MS Word (.docx) documents.
Supports both SRAIS and PII evaluation notebooks.
"""

import argparse
import json
import os
import re
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Color Palette: Academic / Professional Tech
COLOR_PRIMARY_NAVY = RGBColor(0x1E, 0x3A, 0x8A)      # #1E3A8A - Heading 1 & Title
COLOR_SECONDARY_BLUE = RGBColor(0x25, 0x63, 0xEB)    # #2563EB - Heading 2
COLOR_SLATE_HEADER = RGBColor(0x33, 0x41, 0x55)      # #334155 - Heading 3
COLOR_BODY_TEXT = RGBColor(0x0F, 0x17, 0x2A)         # #0F172A - Body text
COLOR_MUTED_GRAY = RGBColor(0x64, 0x74, 0x8B)        # #64748B - Metadata, captions
COLOR_CODE_TEXT = RGBColor(0x0F, 0x17, 0x2A)         # #0F172A - Code text

HEX_TABLE_HEADER_BG = "1E3A8A"
HEX_TABLE_ALT_BG = "F8FAFC"
HEX_CODE_BG = "F8FAFC"
HEX_CODE_BORDER = "CBD5E1"
HEX_CALLOUT_BG = "EFF6FF"
HEX_CALLOUT_BORDER = "3B82F6"


def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Set cell padding (in twips: 20 twips = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_callout_border(cell, border_hex=HEX_CALLOUT_BORDER, size="24"):
    """Set a thick left border on a cell for callouts."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{border_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_box_borders(cell, border_hex=HEX_CODE_BORDER, size="4"):
    """Set light borders around all 4 sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="{border_hex}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{border_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{border_hex}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="{border_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def parse_inline_markdown(paragraph, text):
    """
    Parses inline markdown tokens (**bold**, *italic*, `code`) and creates runs.
    """
    token_pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
    parts = token_pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_BODY_TEXT
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2 and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_BODY_TEXT
        elif part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(label)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = COLOR_SECONDARY_BLUE
                run.underline = True
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_BODY_TEXT


def add_markdown_table(doc, table_lines):
    """
    Parses markdown table lines and renders a styled docx Table.
    """
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line or not line.startswith('|'):
            continue
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    
    if not rows:
        return
    
    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    for r_idx, row_data in enumerate(rows):
        is_header = (r_idx == 0)
        row = table.rows[r_idx]
        
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if is_header:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        
        for c_idx in range(num_cols):
            cell = row.cells[c_idx]
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            
            if is_header:
                set_cell_background(cell, HEX_TABLE_HEADER_BG)
                set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
            else:
                bg = HEX_TABLE_ALT_BG if r_idx % 2 == 1 else "FFFFFF"
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            
            set_box_borders(cell, border_hex="E2E8F0", size="4")
            
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            
            if is_header:
                run = p.add_run(cell_text)
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                parse_inline_markdown(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9.5)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)


def add_code_block(doc, code_str, title=None, language="python"):
    """
    Adds a shaded, monospaced code listing inside a single-cell container table.
    """
    if title:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(3)
        run_label = p_title.add_run(title)
        run_label.bold = True
        run_label.font.name = 'Calibri'
        run_label.font.size = Pt(9.5)
        run_label.font.color.rgb = COLOR_SECONDARY_BLUE

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, HEX_CODE_BG)
    set_box_borders(cell, border_hex=HEX_CODE_BORDER, size="6")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    cell.text = ""
    lines = code_str.split("\n")
    for idx, line in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        
        run = p.add_run(line if line else " ")
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        
        stripped = line.strip()
        if stripped.startswith("#"):
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            run.italic = True
        elif stripped.startswith('"""') or stripped.startswith("'''"):
            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
        else:
            run.font.color.rgb = COLOR_CODE_TEXT

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(4)
    p_spacer.paragraph_format.space_after = Pt(4)


def add_callout_box(doc, text_lines):
    """
    Renders an academic callout box with a prominent left accent border.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, HEX_CALLOUT_BG)
    set_callout_border(cell, border_hex=HEX_CALLOUT_BORDER, size="24")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    cell.text = ""
    for idx, line in enumerate(text_lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        parse_inline_markdown(p, line.strip('> ').strip())
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(4)
    p_spacer.paragraph_format.space_after = Pt(4)


def process_markdown_cell(doc, source_lines):
    """
    Processes a list of markdown lines from a notebook cell into formatted Word elements.
    """
    i = 0
    total = len(source_lines)
    
    while i < total:
        line = source_lines[i].rstrip("\r\n")
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
            
        if re.match(r'^(?:---|\*\*\*|___)$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = [line]
            i += 1
            while i < total:
                next_line = source_lines[i].rstrip("\r\n").strip()
                if next_line.startswith('|'):
                    table_lines.append(next_line)
                    i += 1
                else:
                    break
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith('>'):
            callout_lines = [line]
            i += 1
            while i < total:
                next_line = source_lines[i].rstrip("\r\n").strip()
                if next_line.startswith('>'):
                    callout_lines.append(next_line)
                    i += 1
                else:
                    break
            add_callout_box(doc, callout_lines)
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Calibri Light'
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Calibri Light'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = COLOR_SECONDARY_BLUE
            i += 1
            continue

        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = COLOR_SLATE_HEADER
            i += 1
            continue

        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = True
            run.italic = True
            run.font.color.rgb = COLOR_SLATE_HEADER
            i += 1
            continue

        bullet_match = re.match(r'^[\*\-•]\s+(.*)$', stripped)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_markdown(p, bullet_match.group(1))
            i += 1
            continue

        num_match = re.match(r'^\d+\.\s+(.*)$', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_markdown(p, num_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        parse_inline_markdown(p, stripped)
        i += 1


def build_word_document(notebook_path, output_docx_path, suite_title=None, subtitle=None, headers=None, values=None):
    print(f"Reading notebook from: {notebook_path}")
    with open(notebook_path, "r", encoding="utf-8") as f:
        nb = json.load(f)

    doc = docx.Document()

    # Determine Suite Metadata Defaults
    is_pii = "pii" in notebook_path.lower()
    default_title = "Atticus PII MCP Evaluation Suite" if is_pii else "Atticus SRAIS MCP Evaluation Suite"
    default_sub = (
        "Empirical Validation of Model Context Protocol (MCP) Personally Identifiable Information Detection, Proximity Anchoring & Local Privacy Gating"
        if is_pii else
        "Empirical Validation of Model Context Protocol (MCP) Harm-Scanning Invariants, Deobfuscation Pipeline & Risk Stratification Boundaries"
    )
    doc_title = suite_title or default_title
    doc_sub = subtitle or default_sub
    
    meta_headers = headers or ["Framework Version", "Date", "Evaluation Scope", "Compliance Targets"]
    default_vals = (
        ["v2.0.0 (MCP SSE)", "September 2026", "20-Case Multi-ID Corpus", "GDPR / NIST 800-122 / HIPAA"]
        if is_pii else
        ["v2.0.0 (MCP SSE)", "September 2026", "20-Case Adversarial Corpus", "EU AI Act / NIST AI RMF 1.0"]
    )
    meta_values = values or default_vals

    # Configure Margins: 1 inch (72 pt) all around
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header setup
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"{doc_title} | Academic Review")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED_GRAY
        
        # Footer setup
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("Safe & Responsible AI (SRAI) & Privacy Framework — Confidential / Academic Review")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED_GRAY

    # Document Header Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run(doc_title)
    run_title.font.name = 'Calibri Light'
    run_title.font.size = Pt(26)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY_NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(14)
    run_sub = subtitle_p.add_run(doc_sub)
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY_BLUE

    # Metadata Badges Table
    meta_table = doc.add_table(rows=2, cols=len(meta_headers))
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for c_idx in range(len(meta_headers)):
        h_cell = meta_table.cell(0, c_idx)
        v_cell = meta_table.cell(1, c_idx)
        
        set_cell_background(h_cell, "F1F5F9")
        set_cell_background(v_cell, "FFFFFF")
        set_box_borders(h_cell, border_hex="CBD5E1", size="4")
        set_box_borders(v_cell, border_hex="CBD5E1", size="4")
        set_cell_margins(h_cell, top=80, bottom=80, left=100, right=100)
        set_cell_margins(v_cell, top=80, bottom=80, left=100, right=100)
        
        hp = h_cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hrun = hp.add_run(meta_headers[c_idx])
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.bold = True
        hrun.font.color.rgb = COLOR_SLATE_HEADER
        
        vp = v_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vrun = vp.add_run(meta_values[c_idx])
        vrun.font.name = 'Calibri'
        vrun.font.size = Pt(9.5)
        vrun.font.color.rgb = COLOR_PRIMARY_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Process all notebook cells
    code_cell_counter = 1
    total_cells = len(nb["cells"])
    print(f"Processing {total_cells} notebook cells...")

    for c_idx, cell in enumerate(nb["cells"]):
        cell_type = cell.get("cell_type", "")
        cell_id = cell.get("id", f"cell-{c_idx+1}")
        source_lines = cell.get("source", [])
        
        if cell_type == "markdown":
            process_markdown_cell(doc, source_lines)
            
        elif cell_type == "code":
            code_text = "".join(source_lines)
            
            cell_title = f"Code Listing {code_cell_counter}: [{cell_id}]"
            first_lines = [l.strip() for l in source_lines if l.strip()]
            for l in first_lines[:5]:
                if l.startswith("# Cell") or l.startswith("# =="):
                    continue
                if l.startswith("#") and len(l) > 3 and not l.startswith("# =="):
                    cleaned = l.strip("# =").strip()
                    if cleaned and not cleaned.startswith("==="):
                        cell_title += f" — {cleaned}"
                        break
                        
            add_code_block(doc, code_text, title=cell_title, language="python")
            
            outputs = cell.get("outputs", [])
            for out in outputs:
                out_type = out.get("output_type", "")
                out_text = ""
                if out_type == "stream":
                    out_text = "".join(out.get("text", []))
                elif out_type in ("execute_result", "display_data"):
                    data = out.get("data", {})
                    if "text/plain" in data:
                        out_text = "".join(data["text/plain"])
                
                if out_text.strip():
                    p_out_label = doc.add_paragraph()
                    p_out_label.paragraph_format.space_before = Pt(4)
                    p_out_label.paragraph_format.space_after = Pt(2)
                    r_out = p_out_label.add_run("Output:")
                    r_out.font.name = 'Calibri'
                    r_out.font.size = Pt(8.5)
                    r_out.font.color.rgb = COLOR_MUTED_GRAY
                    r_out.italic = True
                    
                    add_code_block(doc, out_text.strip(), title=None)

            code_cell_counter += 1

    print(f"Saving MS Word document to: {output_docx_path}")
    doc.save(output_docx_path)
    file_size = os.path.getsize(output_docx_path)
    print(f"[SUCCESS] Document successfully created: {output_docx_path} ({file_size:,} bytes)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Jupyter Notebooks to Academic MS Word (.docx) Documents.")
    parser.add_argument("--notebook", "-n", default=r"c:\JDAI\GitHub\atticus\evals()\PII-MCP-Evaluations.ipynb", help="Path to input .ipynb file")
    parser.add_argument("--output", "-o", default=r"c:\JDAI\GitHub\atticus\evals()\PII-MCP-Evaluations.docx", help="Path to output .docx file")
    parser.add_argument("--title", "-t", default=None, help="Document title override")
    args = parser.parse_args()

    build_word_document(args.notebook, args.output, suite_title=args.title)
